'''
Created on Jul 7, 2017

@author: dylan.hematillake
'''
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib.ticker import FormatStrFormatter
import numpy as np
import time as t
from win32com import client
from docx import Document 
from docx.shared import Inches

'''Import the data in csv format to the program from the specified folder on the desktop, can be changed'''
def input_data():
    name = input("Enter file name with extension, ex data.csv:   ")
    
    '''Change the destination as required'''
    filename = "C:/Users/dylan.hematilake/Desktop/Vibration Data/"+name
    print(filename)
    vibration_data = pd.read_csv(filename,sep=";")
    print("Author: Dylan Hematillake")
    print("Loaded Data:")
    print("-------------------------------------------------------------------------------------------------------------")
    print(vibration_data)
    print("-------------------------------------------------------------------------------------------------------------")
    return vibration_data

word = client.Dispatch("Word.Application")

'''allows for the graph to be printed if the file is saved in Documents'''
def printWordDocument(filename):

    word.Documents.Open(filename)
    word.ActiveDocument.PrintOut()
    t.sleep(5)
    word.ActiveDocument.Close()

'''print the fig, saves as word doc'''
def PrintMethod(printStatus,fig):
                        
        if printStatus == 'yes':
            
            '''File name extension must be docx'''
            
            filename = input("Name your file, ex 'Plot.docx'")
            name,ext = filename.split('.')
            fig.savefig(name+'.png')
            document = Document()
            document.add_heading('Plot',0)
            document.add_picture(name+'.png',width=Inches(7))
            '''file must be saved to Documents in order to print, change user as required'''
            document.save('C:/Users/dylan.hematilake/Documents/'+filename)
            printWordDocument(filename)
        word.Quit()
          
        if printStatus == 'no':
                plt.show() 

'''Labels plots for the method below'''
def autolabel(rects,ax):
        for rect in rects:
            height = rect.get_height()
            ax.text(rect.get_x() + rect.get_width()/2., 1.05*height/2,
                    '%.2f' % height,
                    ha='center', va='bottom')

'''the main method to extract and categorize the csv data and produce the graphs'''    
def process_data():
    vibration_data = input_data()
    Product = vibration_data.iloc[:,5]
    ahv = vibration_data.iloc[:,12]
    eav = vibration_data.iloc[:,15]
    
    '''neccessary formatting due to the output from the sensor software'''
    def format():
        ahvBH,units1 = ahv[0].split(' ')
        h,m = eav[0].split(':')
        eavBH = int(h)*3600+int(m)*60
        eavHMFormat = []
        for i in range(0,len(ahv)):
            hold = []
            hold2 = []
            tool,product = Product[i].split(':')
            num, unit =ahv[i].split(' ')
            num = ((float(ahvBH)-float(num))/float(ahvBH))*100
            hrs,mins= eav[i].split(':')
            timeInSecs = int(hrs)*3600+int(mins)*60-eavBH
            hold2.append(timeInSecs/60)
            
            if timeInSecs < 0:
                timeInSecs = 86400 - timeInSecs
            #eavHMFormat.append(t.strftime('%H:%M:%S',t.gmtime(timeInSecs)))
            eav.iloc[i]=hold2[0]
            hold.append(num)
            ahv.iloc[i]=hold[0]
        return eav,ahv
    
    def plotGraph(ax,ax2):
            rects1 = ax.bar(ind,ahv.iloc[1:],width,color='#ff9966',align='edge')
            ax.set_ylabel('AHV Improvement (%)')
            ax.set_xlabel("Tool Used")
            ax.set_title("Improvement to AHV")
            ax.set_xticks(ind+width/2)
            ax.set_xticklabels(Product.iloc[1:])
            ax.yaxis.set_major_formatter(FormatStrFormatter('%.2f'))
        
            #eav plot
            rects2 = ax2.bar(ind,eav.iloc[1:],width,color='#7f7fff',align="edge")
            ax2.set_ylabel('EAV Time Added (Mins)')
            ax2.set_xlabel("Tool Used")
            ax2.set_title("Improvement to EAV (Mins)")
            ax2.set_xticks(ind+width/2)
            ax2.set_xticklabels(Product.iloc[1:])
            ax2.tick_params(axis='x',pad=15)
    
            autolabel(rects1,ax)
            autolabel(rects2,ax2)
            
    testType = input("Are you performing a preliminary test? 'yes' or 'no':  ")
    
    if testType == 'yes':
        format()
        #plot data
        N = len(Product)-1
        ind = np.arange(N)
        width = 0.35
        
        fig,(ax,ax2) = plt.subplots(2,sharex=True,figsize=(10,10))
        plotGraph(ax, ax2)
        printStatus = input("Do you want to print a hard copy? 'yes' or 'no':  ")
        PrintMethod(printStatus, fig) 
            
            
    elif testType == 'no':
        format()
        numTests = input("Are you inputting data from multiple tools? Answer 'yes' or 'no':  ")

        while numTests == 'yes':
            toolList = []
            toolList.append(input("Input the tools you are using one at a time. You will be prompted per additional entry;  "))
            check = input("Are you inputting another tool? 'yes' or 'no':  ")
            if check != 'yes':
                break
        jhAHV,hdAHV,grdrAHV,hmrdrillAHV = [[],[],[],[]]
        jhEAV,hdEAV,grdrEAV,hmrdrillEAV = [[],[],[],[]]
        
        '''
            The following will cause a bug if the user does not fill the 'Task' box properly.
            It must read exactly Tool:Product
            example: Jackhammer:S10VIB
        '''
      
        if numTests == "yes":
            for i in range(0,len(Product)):
                tool,product = Product[i].split(':')
                if tool == "Jackhammer":
                    jhAHV.append(ahv.iloc[i])
                    jhEAV.append(eav.iloc[i])
                elif tool == 'Hammer Drill':
                    hmrdrillAHV.append(ahv.iloc[i])
                    hmrdrillEAV.append(eav.iloc[i])
                elif tool == 'Grinder':
                    grdrAHV.append(ahv.iloc[i])
                    grdrEAV.append(eav.iloc[i])
                elif tool == 'Hand Drill':
                    hdAHV.append(ahv.iloc[i])
                    hdEAV.append(eav.iloc[i])
        #plot data
        N = len(Product)-1
        ind = np.arange(N)
        width = 0.35
        
        fig,(ax,ax2) = plt.subplots(2,sharex=True, figsize=(10,10))
        plotGraph(ax, ax2)
        printStatus = input("Do you want to print a hard copy? 'yes' or 'no':  ")
        PrintMethod(printStatus, fig)
        
    else:
        print("Invalid Response")
    
def main():
    process_data()
    
if __name__ == '__main__':
    main()

